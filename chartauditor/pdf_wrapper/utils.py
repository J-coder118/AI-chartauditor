from pdfminer.layout import LTTextContainer, LTChar
from django.template.loader import render_to_string
from channels.layers import get_channel_layer
from chartauditor.pdf_wrapper.models import ChartChecker
from asgiref.sync import async_to_sync
from django.http import HttpResponse
from nltk.corpus import stopwords
from django.conf import settings
from celery import shared_task
from nltk import word_tokenize
from weasyprint import HTML
from openai import OpenAI
from docx import Document
from io import BytesIO
import zipfile
import PyPDF2
import string
import time
import re


def send_progress(progress, user_id):
    channel_layer = get_channel_layer()
    async_to_sync(channel_layer.group_send)(
        str(user_id),
        {
            'type': 'chat.message',
            'message': progress,
        }
    )


def make_chunks(data):
    clean_text = ' '.join(data)
    clean_text_len = len(clean_text)
    chunk_size = 28000
    num_chunks = (clean_text_len + chunk_size - 1) // chunk_size
    data_dict = {}

    for i in range(num_chunks):
        start_index = i * chunk_size
        end_index = min((i + 1) * chunk_size, clean_text_len)
        chunk_key = f'chunk{i + 1}'
        data_dict[chunk_key] = clean_text[start_index:end_index]
    return data_dict


def update_chart_obj_status(chart_obj, user):
    chart_obj.is_payment_done = True
    data_dict = make_chunks(chart_obj.clean_chart)
    chart_obj.clean_chart = data_dict
    chart_obj.save()
    user.character_limit = user.character_limit + chart_obj.character_count
    user.save()

    assistant_id = None
    if chart_obj.is_report_emailed:
        assistant_id = settings.FLORIDA_COMPLIANCE_ASSISTANT_ID
    elif chart_obj.is_state_compliance:
        if chart_obj.state_compliance == 'California':
            assistant_id = settings.CALIFORNIA_COMPLIANCE_ASSISTANT_ID
        else:
            assistant_id = settings.FLORIDA_COMPLIANCE_ASSISTANT_ID
    elif chart_obj.is_insurance_compliance:
        if chart_obj.state_compliance == 'Aetna':
            assistant_id = settings.AETNA_ASSISTANT_ID
        else:
            assistant_id = settings.CIGNA_ASSISTANT_ID
    elif chart_obj.is_CARF_compliance:
        assistant_id = settings.CARF_COMPLIANCE_ASSISTANT_ID
    elif chart_obj.is_marked_cover_letter:
        assistant_id = settings.FLORIDA_COMPLIANCE_ASSISTANT_ID
    elif chart_obj.is_commission_compliance:
        assistant_id = settings.JOINT_COMMISSION_ASSISTANT_ID

    create_thread_and_run.delay(chart_obj.clean_chart, chart_obj.id, assistant_id)


def tokenize_data(data):
    tokenized_list = []
    [tokenized_list.extend(word_tokenize(item)) for item in data]
    return tokenized_list


def pdf_reader(file):
    extracted_text = ""
    page_count = 0
    with BytesIO(file.read()) as file:
        reader = PyPDF2.PdfReader(file)
        page_count = len(reader.pages)
        for page_num in range(page_count):
            page = reader.pages[page_num]
            extracted_text += page.extract_text()
    return extracted_text, page_count


def process_zip_file(file, combined_text, total_page_count):
    with zipfile.ZipFile(file, 'r') as z:
        for file_info in z.infolist():
            if file_info.is_dir():
                # Skip directories
                continue
            if file_info.filename.lower().endswith('.pdf'):
                with z.open(file_info.filename) as pdf_file:
                    extracted_text, page_count = pdf_reader(pdf_file)
                    combined_text.append(extracted_text)
                    total_page_count += page_count
            elif file_info.filename.lower().endswith('.zip'):
                with z.open(file_info.filename) as nested_zip_file:
                    nested_zip_contents = BytesIO(nested_zip_file.read())
                    total_page_count = process_zip_file(nested_zip_contents, combined_text, total_page_count)
    return total_page_count


def file_processing(file, user_id):
    progress = '11-70'
    combined_text = []
    total_page_count = 0

    if file.name.lower().endswith('.pdf'):
        extracted_text, page_count = pdf_reader(file)
        combined_text.append(extracted_text)
        total_page_count += page_count
    elif file.name.lower().endswith('.zip'):
        process_zip_file(file, combined_text, total_page_count)

    send_progress(progress, user_id)
    return combined_text, total_page_count


# def file_processing(file, user_id):
#     progress = '11-70'
#     combined_text = []
#     total_page_count = 0
#     if file.name.endswith('.zip'):
#         with zipfile.ZipFile(file, 'r') as z:
#             for pdf_filename in z.namelist():
#                 if pdf_filename.lower().endswith('.pdf'):
#                     with z.open(pdf_filename) as pdf_file:
#                         extracted_text, page_count = pdf_reader(pdf_file)
#                         combined_text.append(extracted_text)
#                         total_page_count = total_page_count + page_count
#         send_progress(progress, user_id)
#
#     elif file.name.endswith('.pdf'):
#         extracted_text, page_count = pdf_reader(file)
#         combined_text.append(extracted_text)
#         total_page_count = total_page_count + page_count
#         send_progress(progress, user_id)
#     return combined_text, total_page_count


def cost_calculation(clean_text):
    clean_text_len = len(clean_text)
    chart_cost = (clean_text_len / 2800) * 5
    rounded_chart_cost = round(chart_cost, 3)
    return rounded_chart_cost


def remove_stopwords(tokens_without_punctuation):
    tokenized_without_stopwords = []
    for token in tokens_without_punctuation:
        if not token in stopwords.words('english'):
            tokenized_without_stopwords.append(token)
    return tokenized_without_stopwords


def remove_punctuation_and_stopwords(tokenize_text):
    regex = re.compile('[%s]' % re.escape(string.punctuation))
    tokens_without_punctuation = []
    for token in tokenize_text:
        new_token = regex.sub(u'', token)
        if not new_token == u'':
            tokens_without_punctuation.append(new_token)
    clean_tokens = remove_stopwords(tokens_without_punctuation)
    return clean_tokens


def get_user_inputs(request):
    data = [request.POST.get('first_name', '').lower(), request.POST.get('last_name', '').lower(),
            request.POST.get('dob', '').lower(), request.POST.get('mr_number', '').lower(),
            request.POST.get('ss_number', '').lower(), request.POST.get('driving_license', '').lower(),
            request.POST.get('address', '').lower(), request.POST.get('credit_card', '').lower()]
    tokenized_list = tokenize_data(data)
    clean_tokens = remove_punctuation_and_stopwords(tokenized_list)
    return clean_tokens


def remove_user_inputs(chart_obj, tokenized_user_inputs_no_punctuation):
    clean_tokens = []
    for token in chart_obj.clean_chart:
        if token not in tokenized_user_inputs_no_punctuation:
            clean_tokens.append(token)
    data_dict = make_chunks(clean_tokens)
    return data_dict


def render_to_pdf_weasy(template_src, request, filename, context_dict={}):
    html = render_to_string(template_src, context_dict)
    pdf = HTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf()
    response = HttpResponse(pdf, content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename={filename}.pdf'
    return response


def convert_to_pdf(doc_buffer, file_name):
    docx_filename = f"{file_name}.docx"
    pdf_filename = f"{file_name}.pdf"
    pass


def generate_docx(file_name, content, btn):
    doc = Document()

    for line in content.split('\n'):
        if line.startswith('#'):
            header_level = line.count('#')
            line = line.lstrip('#').strip()
            doc.add_heading(line, level=header_level)
        else:
            # Add paragraph with regular styling
            doc.add_paragraph(line)

    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    if btn == 'pdf':
        pdf_buffer = convert_to_pdf(doc_buffer, file_name)
        response = HttpResponse(pdf_buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{file_name}.pdf"'
        return response

    response = HttpResponse(
        doc_buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{file_name}.docx"'
    return response


def wait_on_run(run, thread, client):
    while run.status == "queued" or run.status == "in_progress":
        run = client.beta.threads.runs.retrieve(
            thread_id=thread.id,
            run_id=run.id,
        )
        time.sleep(0.5)
    return run


def get_response(thread, client):
    message_list = client.beta.threads.messages.list(thread_id=thread.id, order="asc")
    for m in message_list:
        print("for loop print")
    return m.content[0].text.value


@shared_task()
def create_thread_and_run(data_dict, obj_id, assistant_id):
    client = OpenAI(api_key=settings.OPENAI_API_KEY)
    thread = client.beta.threads.create()
    for chunk_key, chunk_text in data_dict.items():
        client.beta.threads.messages.create(
            thread_id=thread.id, role="user", content=chunk_text
        )
    run_obj = client.beta.threads.runs.create(
        thread_id=thread.id,
        assistant_id=assistant_id,
    )
    wait_on_run(run_obj, thread, client)
    response = get_response(thread, client)

    score_pattern = r"Cumulative \"ChartAuditor\.com Score\": (\d+/\d+)|Overall \*\*ChartAuditor\.com Score: (\d+/\d+)\.|\*\*ChartAuditor\.com Score: (\d+/\d+)\.|Cumulative Score \(ChartAuditor\.com Score\): (\d+/\d+)|\*\*Overall ChartAuditor\.com Score: (\d+/\d+)\.|\*\*Cumulative ChartAuditor\.com Score: (\d+/\d+)\."
    cumulative_score = re.search(score_pattern, response)

    instance = ChartChecker.objects.get(id=obj_id)
    instance.chart_response = response
    cumulative_score_value = ''

    if cumulative_score:
        # Check which group matched
        for score_group in cumulative_score.groups():
            if score_group:
                cumulative_score_value = score_group
                break
    instance.score = cumulative_score_value
    instance.save()


def text_extraction(element):
    line_text = element.get_text()
    line_formats = []
    for text_line in element:
        if isinstance(text_line, LTTextContainer):
            for character in text_line:
                if isinstance(character, LTChar):
                    line_formats.append(character.fontname)
                    line_formats.append(character.size)
    format_per_line = list(set(line_formats))
    return line_text, format_per_line


# def pdf_reader(file):
#     temp_file = tempfile.NamedTemporaryFile(delete=True)
#     temp_file.write(file.read())
#
#     pdf_path = temp_file.name
#     pdfFileObj = open(pdf_path, 'rb')
#
#     text_per_page = {}
#     page_count = 0
#     for pagenum, page in enumerate(extract_pages(pdf_path)):
#         page_count += 1
#
#         page_text = []
#         line_format = []
#         page_content = []
#
#         for element in page:
#             if isinstance(element, LTTextContainer):
#                 (line_text, format_per_line) = text_extraction(element)
#                 page_text.append(line_text)
#                 line_format.append(format_per_line)
#                 page_content.append(line_text)
#
#         dctkey = 'Page_' + str(pagenum)
#         text_per_page[dctkey] = [page_text, line_format, page_content]
#
#     pdfFileObj.close()
#     result = ''.join(''.join(page[2]) for page in text_per_page.values())
#
#     return result, page_count
